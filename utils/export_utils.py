"""
Export utilities for University Management System
Supports export to Excel, Word, and PDF formats
"""

import pandas as pd
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import io
import os
from datetime import datetime

def export_students_to_excel(students_data):
    """Export students data to Excel format"""
    try:
        # Create DataFrame
        df = pd.DataFrame(students_data)
        
        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Students', index=False)
            
            # Get the workbook and worksheet
            workbook = writer.book
            worksheet = writer.sheets['Students']
            
            # Auto-adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        print(f"Error exporting to Excel: {e}")
        return None

def export_students_to_pdf(students_data):
    """Export students data to PDF format"""
    try:
        # Create PDF in memory
        output = io.BytesIO()
        doc = SimpleDocTemplate(output, pagesize=A4)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Content
        story = []
        
        # Title
        title = Paragraph("Students Report", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Date
        date_str = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        date_para = Paragraph(date_str, styles['Normal'])
        story.append(date_para)
        story.append(Spacer(1, 12))
        
        # Table data
        if students_data:
            # Headers
            headers = list(students_data[0].keys())
            table_data = [headers]
            
            # Data rows
            for student in students_data:
                row = [str(student.get(key, '')) for key in headers]
                table_data.append(row)
            
            # Create table
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
        else:
            no_data = Paragraph("No student data available", styles['Normal'])
            story.append(no_data)
        
        # Build PDF
        doc.build(story)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        print(f"Error exporting to PDF: {e}")
        return None

def export_students_to_word(students_data):
    """Export students data to Word format"""
    try:
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Students Report', 0)
        title.alignment = 1  # Center alignment
        
        # Date
        date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = 1
        
        # Add space
        doc.add_paragraph()
        
        if students_data:
            # Create table
            headers = list(students_data[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = str(header)
                hdr_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for student in students_data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(student.get(header, ''))
        else:
            doc.add_paragraph("No student data available")
        
        # Save to memory
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        print(f"Error exporting to Word: {e}")
        return None

def export_courses_to_excel(courses_data):
    """Export courses data to Excel format"""
    try:
        df = pd.DataFrame(courses_data)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Courses', index=False)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        print(f"Error exporting courses to Excel: {e}")
        return None

def export_teachers_to_excel(teachers_data):
    """Export teachers data to Excel format"""
    try:
        df = pd.DataFrame(teachers_data)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Teachers', index=False)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        print(f"Error exporting teachers to Excel: {e}")
        return None

